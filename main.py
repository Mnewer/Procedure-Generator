"""Procedure Generator"""

"""Import Modules"""
#.docx files:
import docx
#.PDF files:
import PyPDF2
#.PDF files:
from fpdf import FPDF
#GUI:
import tkinter

"""Test area:"""

"""Open template"""
#Template in .docx or .PDF
importFileFromLocation = ()


#.docx
fileLocation = ()
template = docx.Document('Test.docx')
fullText = []
for paragraph in template.paragraphs:
    fullText.append(paragraph.text)

#Need to convert fullList to a string for fpdf to manage it:
def listToString(s):  
    str1 = " " 
    return (str1.join(s)) 
        
s = (fullText)
print(listToString(s))
fullTextstr = listToString
#print(fullText)
#print(len(template.paragraphs))


"""Select sections from template"""
#Select all 


"""Make PDF from selected text from docx"""


"""Make new PDF file with pyPDF2, only generates a new blank PDF"""
# """Export newly made procedure to .PDF or .docx"""
# exportFileToLocation = ()
# """Test area:"""
"""fpdf"""
# newPDF = FPDF()
# newPDF.add_page()
# newPDF.set_font('Arial', size = 15)
# #textFrom = open()
# newPDF.cell(w = 0, h = 0, txt='Test.docx', border=0, ln=0, align='', fill=0, link='')
# newPDF.output('test.pdf')

# from PyPDF2 import PdfFileWriter as w
# """PDF Writer:"""


# """Output PDF:"""
# outputPdf = w()
# file = open('Test.docx', 'wb')
# for i in range(2):
#     outputPdf.getObject
# outputPdf.write(file)
# file.close()



"""Saving multiple files. fileName+1 to avoid overwriting.
fileName = 'Title, type of procedure and number'
typeProcedure = 'Name based on selected options'
optionList = ['Option1', 'Option2', 'Option3']
"""

#.docx



#.pdf









"""Notes:
Single line comment:
Ctrl + K, Ctrl + C
Single line uncomment:
Ctrl + K, Ctrl + U
"""